"""
Document Import/Export Services for Legenddary Platform
Handles Word import, multiple export formats, AI content analysis
"""
import io
import re
import base64
from typing import List, Dict, Tuple
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as pdf_canvas


def parse_docx(file_content: bytes) -> List[Dict]:
    """Parse a Word document and extract chapters/sections"""
    doc = Document(io.BytesIO(file_content))
    
    sections = []
    current_section = {"title": "Imported Content", "type": "chapter", "content": ""}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            current_section["content"] += "<p><br/></p>"
            continue
        
        style = para.style.name.lower() if para.style else ""
        
        # Detect headers/chapters
        if "heading 1" in style or is_chapter_header(text):
            if current_section["content"].strip():
                sections.append(current_section)
            current_section = {
                "title": text,
                "type": detect_section_type(text),
                "content": ""
            }
        elif "heading 2" in style:
            current_section["content"] += f"<h2>{text}</h2>"
        elif "heading 3" in style:
            current_section["content"] += f"<h3>{text}</h3>"
        else:
            # Check for bold/italic
            formatted_text = text
            for run in para.runs:
                if run.bold:
                    formatted_text = formatted_text.replace(run.text, f"<strong>{run.text}</strong>")
                if run.italic:
                    formatted_text = formatted_text.replace(run.text, f"<em>{run.text}</em>")
            current_section["content"] += f"<p>{formatted_text}</p>"
    
    if current_section["content"].strip():
        sections.append(current_section)
    
    return sections if sections else [{"title": "Chapter 1", "type": "chapter", "content": "<p>No content found</p>"}]


def is_chapter_header(text: str) -> bool:
    """Detect if text is a chapter header"""
    patterns = [
        r'^chapter\s+\d+',
        r'^part\s+\d+',
        r'^section\s+\d+',
        r'^prologue',
        r'^epilogue',
        r'^introduction',
        r'^preface',
        r'^dedication',
        r'^acknowledgment',
        r'^appendix',
    ]
    text_lower = text.lower().strip()
    return any(re.match(p, text_lower) for p in patterns)


def detect_section_type(title: str) -> str:
    """Detect section type from title"""
    title_lower = title.lower()
    if "prologue" in title_lower:
        return "prologue"
    elif "epilogue" in title_lower:
        return "epilogue"
    elif "preface" in title_lower:
        return "preface"
    elif "introduction" in title_lower:
        return "introduction"
    elif "dedication" in title_lower:
        return "dedication"
    elif "acknowledgment" in title_lower:
        return "acknowledgments"
    elif "appendix" in title_lower:
        return "afterword"
    else:
        return "chapter"


def export_to_html(book: dict, chapters: list) -> str:
    """Export book to HTML format"""
    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{book['title']}</title>
    <style>
        body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 40px 20px; line-height: 1.8; }}
        h1 {{ text-align: center; margin-bottom: 40px; }}
        h2 {{ margin-top: 60px; page-break-before: always; }}
        p {{ text-indent: 1.5em; margin: 0.5em 0; }}
        .chapter {{ margin-bottom: 60px; }}
        .title-page {{ text-align: center; margin-bottom: 100px; }}
        .title-page h1 {{ font-size: 2.5em; }}
        .title-page .description {{ color: #666; font-style: italic; }}
    </style>
</head>
<body>
    <div class="title-page">
        <h1>{book['title']}</h1>
        <p class="description">{book.get('description', '')}</p>
    </div>
"""
    for chapter in chapters:
        html += f"""
    <div class="chapter">
        <h2>{chapter['title']}</h2>
        {chapter.get('content', '')}
    </div>
"""
    html += "</body></html>"
    return html


def export_to_txt(book: dict, chapters: list) -> str:
    """Export book to plain text format"""
    text = f"{book['title'].upper()}\n{'=' * len(book['title'])}\n\n"
    if book.get('description'):
        text += f"{book['description']}\n\n"
    text += "-" * 50 + "\n\n"
    
    for chapter in chapters:
        text += f"\n{chapter['title'].upper()}\n{'-' * len(chapter['title'])}\n\n"
        content = chapter.get('content', '')
        # Strip HTML tags
        content = re.sub(r'<[^>]+>', '', content)
        content = re.sub(r'&nbsp;', ' ', content)
        content = re.sub(r'&amp;', '&', content)
        content = re.sub(r'&lt;', '<', content)
        content = re.sub(r'&gt;', '>', content)
        content = re.sub(r'\s+', ' ', content).strip()
        
        # Word wrap at 80 chars
        words = content.split()
        line = ""
        for word in words:
            if len(line) + len(word) + 1 > 80:
                text += line + "\n"
                line = word
            else:
                line = line + " " + word if line else word
        if line:
            text += line + "\n"
        text += "\n"
    
    return text


def export_to_docx(book: dict, chapters: list) -> bytes:
    """Export book to Word document format"""
    doc = Document()
    
    # Title
    title = doc.add_heading(book['title'], 0)
    title.alignment = 1  # Center
    
    if book.get('description'):
        p = doc.add_paragraph(book['description'])
        p.alignment = 1
    
    doc.add_page_break()
    
    for chapter in chapters:
        doc.add_heading(chapter['title'], 1)
        
        content = chapter.get('content', '')
        # Parse HTML content
        content = re.sub(r'<h2>(.*?)</h2>', r'\n## \1\n', content)
        content = re.sub(r'<h3>(.*?)</h3>', r'\n### \1\n', content)
        content = re.sub(r'<strong>(.*?)</strong>', r'\1', content)
        content = re.sub(r'<em>(.*?)</em>', r'\1', content)
        content = re.sub(r'<p>(.*?)</p>', r'\1\n', content)
        content = re.sub(r'<[^>]+>', '', content)
        content = re.sub(r'\s+', ' ', content).strip()
        
        paragraphs = content.split('\n')
        for para in paragraphs:
            para = para.strip()
            if para.startswith('## '):
                doc.add_heading(para[3:], 2)
            elif para.startswith('### '):
                doc.add_heading(para[4:], 3)
            elif para:
                doc.add_paragraph(para)
        
        doc.add_page_break()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def analyze_content_structure(content: str) -> Dict:
    """AI-powered content analysis for smart recommendations"""
    analysis = {
        "detected_headers": [],
        "potential_chapters": [],
        "suggestions": [],
        "word_count": len(content.split()),
        "paragraph_count": content.count('<p>') or content.count('\n\n') + 1
    }
    
    # Detect potential chapter breaks
    chapter_patterns = [
        (r'chapter\s+(\d+|[ivxlc]+)', 'Chapter'),
        (r'part\s+(\d+|[ivxlc]+)', 'Part'),
        (r'^#{1,3}\s+(.+)$', 'Heading'),
    ]
    
    lines = content.split('\n')
    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        for pattern, type_name in chapter_patterns:
            if re.search(pattern, line_lower, re.IGNORECASE):
                analysis["detected_headers"].append({
                    "line": i,
                    "text": line.strip(),
                    "type": type_name
                })
    
    # Generate suggestions
    if analysis["word_count"] > 5000 and len(analysis["detected_headers"]) == 0:
        analysis["suggestions"].append({
            "type": "structure",
            "message": "Your content is long but has no chapter breaks. Consider adding chapter divisions for better readability."
        })
    
    if analysis["paragraph_count"] < analysis["word_count"] / 500:
        analysis["suggestions"].append({
            "type": "improvement",
            "message": "Your paragraphs seem long. Consider breaking them up for easier reading."
        })
    
    return analysis


def smart_split_content(content: str, split_by: str = "chapter") -> List[Dict]:
    """Intelligently split content into chapters"""
    chapters = []
    
    if split_by == "chapter":
        pattern = r'(chapter\s+\d+[:\.\s]*[^\n]*)'
    elif split_by == "part":
        pattern = r'(part\s+\d+[:\.\s]*[^\n]*)'
    elif split_by == "heading":
        pattern = r'(#{1,3}\s+[^\n]+)'
    else:
        pattern = r'(chapter\s+\d+[:\.\s]*[^\n]*)'
    
    parts = re.split(pattern, content, flags=re.IGNORECASE)
    
    current_chapter = {"title": "Introduction", "content": ""}
    
    for i, part in enumerate(parts):
        if re.match(pattern, part, re.IGNORECASE):
            if current_chapter["content"].strip():
                chapters.append(current_chapter)
            current_chapter = {"title": part.strip(), "content": ""}
        else:
            current_chapter["content"] += part
    
    if current_chapter["content"].strip():
        chapters.append(current_chapter)
    
    return chapters if chapters else [{"title": "Chapter 1", "content": content}]
